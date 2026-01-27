import pandas as pd
from docx import Document
from docx2pdf import convert
import os
from mapeos import TITULO, ESTUDIANTE, READING

def generar_boletin(id_estudiante, ruta_excel, trimestre_a_imprimir):
    # 1. Cargar datos del Excel
    dfnotas = pd.read_excel(ruta_excel, sheet_name='Tablero_notas_Oficial')
    dfcomentarios= pd.read_excel(ruta_excel, sheet_name='Ls_Comments_Oficial')

    # --- CORRECCIÓN CRÍTICA: Limpiar nombres de columnas ---
    # Esto elimina espacios invisibles al inicio/final como 'Trimester1 ' -> 'Trimester1'
    dfnotas.columns = dfnotas.columns.str.strip()

    # 2. NORMALIZACIÓN DE DATOS DEL ESTUDIANTE
    dfnotas['CodigoEstudiante'] = dfnotas['CodigoEstudiante'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
    id_busqueda = str(id_estudiante).split('.')[0].strip()

    # 3. FILTRAR TODAS LAS FILAS DEL ESTUDIANTE
    info = dfnotas[dfnotas['CodigoEstudiante'] == id_busqueda]

    if info.empty:
        print(f"❌ ERROR: No se encontró al estudiante con código: {id_busqueda}")
        return
    
    # 4. DATOS GENERALES
    datos_est = info.iloc[0]
    nombre_alumno = str(datos_est['StudentName']).strip()
    grado_completo = str(datos_est['HR']).strip().replace("\n", "")
    grado_numero = grado_completo[0] if grado_completo else "1"
    
    print(f"✅ Procesando: {nombre_alumno} | Grado: {grado_completo} | Trimestre: {trimestre_a_imprimir}")

    # --- FUNCIÓN DE BÚSQUEDA BLINDADA ---
    def buscar_nota(materia, dominio, num_trimestre):
        columna_nota = f'Trimester{num_trimestre}'
        
        # Verificar si la columna existe después del strip()
        if columna_nota not in info.columns:
            return f"Error: {columna_nota} no encontrada"

        filtro = info[(info['Subject'].str.strip() == materia) & (info['Domain'].str.strip() == dominio)]
        
        if not filtro.empty:
            valor = filtro.iloc[0][columna_nota]
            return str(valor) if pd.notna(valor) else ""
        return ""
    
    def buscar_profesor(materia):
        """
        Busca el nombre del profesor asignado a una materia específica.
        """
        # Filtramos por la columna 'Subject' (quitando espacios en blanco)
        filtro = info[info['Subject'].str.strip() == materia]
        
        if not filtro.empty:
            # Asumiendo que tu columna se llama 'S_Teacher' (según tu etiqueta de AppSheet)
            # Si en tu DataFrame se llama distinto, cambia 'S_Teacher' aquí abajo
            valor = filtro.iloc[0]['S_Teacher']
            return str(valor) if pd.notna(valor) else "Sin profesor"
        
        return "Materia no encontrada"

    # 5. SELECCIÓN DE PLANTILLA
    carpeta_plantillas = "PLANTILLAS"
    plantillas = {
        "1": "Grades1&2template.docx",
        "2": "Grades1&2template.docx",
        "3": "Grades3,4&5template.docx",
        "4": "Grades3,4&5template.docx",
        "5": "Grades3,4&5template.docx"
    }

    nombre_archivo_plantilla = plantillas.get(grado_numero, "Grades1&2template.docx")
    ruta_plantilla = os.path.join(carpeta_plantillas, nombre_archivo_plantilla)
    
    if not os.path.exists(ruta_plantilla):
        print(f"❌ ERROR: No existe la plantilla {ruta_plantilla}")
        return

    doc = Document(ruta_plantilla)

    # 6. MAPEO DE REEMPLAZOS (CABECERA)
    mapeo_reemplazos = {
        TITULO["FINALREPORT"]: f"FINAL REPORT TRIMESTER {trimestre_a_imprimir}",
        ESTUDIANTE["NOMBRE"]: nombre_alumno,
        ESTUDIANTE["GRADO"]: grado_completo,
        ESTUDIANTE["PROFE"]: str(datos_est.get('HR_Teacher', 'N/A')),
        ESTUDIANTE["ID"]: id_busqueda,
        # READING["NOMBREPROFE"]: buscar_nota("Reading", "Foundational Skills", 1)
    }

    # 7. LÓGICA DE TRIMESTRES (ACUMULATIVA Y LIMPIEZA)
    print("============NOTA=====================")
    print(buscar_nota("Reading", "Foundational Skills", 1))
    print("=================================")
    for t in range(1, 4):
        if t <= trimestre_a_imprimir:
            print("ENTRA")
            # Llenar notas reales
            if grado_numero in ["1", "2"]:
                print("ENTRA 2")
                mapeo_reemplazos.update({
                    READING["NOMBREPROFE"]: buscar_profesor("Reading"),
                    READING[f"Literature&Information_T{t}"]: buscar_nota("Reading", "Literature & Information", t)
                    # Agrega más etiquetas según necesites
                })
                print("=================================")
                print(buscar_nota("Reading", "Foundational Skills", 1))
                print("=================================")
        else:
            # Limpiar etiquetas de trimestres futuros (poner en blanco)
            if grado_numero in ["1", "2"]:
                mapeo_reemplazos.update({
                    READING[f"Literature&Information_T{t}"]: ""
                })
            # else:
            #     mapeo_reemplazos.update({
            #         f"<<Reading_Adv_T{t}>>": "",
            #         f"<<Science_T{t}>>": "",
            #     })

    # 8. EJECUTAR REEMPLAZO
    # 8. EJECUTAR REEMPLAZO (Versión Robusta)
    def reemplazar_en_bloque(objeto_con_parrafos):
        for p in objeto_con_parrafos:
            # Primero revisamos si la etiqueta está en el texto completo del párrafo
            for etiqueta, valor in mapeo_reemplazos.items():
                if etiqueta in p.text:
                    # Si la etiqueta existe, la reemplazamos en los runs de forma segura
                    # Intentamos reemplazar en el párrafo completo si los runs fallan
                    for run in p.runs:
                        if etiqueta in run.text:
                            run.text = run.text.replace(etiqueta, str(valor))
                    
                    # Si después de los runs aún queda la etiqueta (porque estaba dividida)
                    # hacemos un reemplazo forzado en el párrafo
                    if etiqueta in p.text:
                        # Nota: Esto puede perder negritas específicas, pero asegura el dato
                        p.text = p.text.replace(etiqueta, str(valor))

    # Ejecutar en párrafos normales
    reemplazar_en_bloque(doc.paragraphs)
    
    # Ejecutar en tablas
    for tabla in doc.tables:
        for fila_t in tabla.rows:
            for celda in fila_t.cells:
                reemplazar_en_bloque(celda.paragraphs)

    # 9. GUARDADO Y CONVERSIÓN
    nombre_limpio = f"Boletin_{grado_completo}_{id_busqueda}".replace(" ", "_").replace("\n", "")
    ruta_word = f"{nombre_limpio}.docx"
    ruta_pdf = f"{nombre_limpio}.pdf"

    doc.save(ruta_word)
    
    try:
        print(f"⏳ Convirtiendo a PDF...")
        convert(ruta_word, ruta_pdf)
        print(f"✅ PDF generado con éxito: {ruta_pdf}")
        if os.path.exists(ruta_pdf):
            os.remove(ruta_word)
    except Exception as e:
        print(f"❌ Error al convertir PDF: {e}. Se conserva el Word.")

# --- EJECUCIÓN ---
generar_boletin(20622, "baseprueba.xlsx", 3)
# generar_boletin(20622, "Data Domains Elementary.xlsx", 1)